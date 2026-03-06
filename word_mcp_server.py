#!/usr/bin/env python3
"""
Word Document MCP Server
========================
A fully-featured MCP server for creating and manipulating .docx files
via Claude Code integration.

Capabilities:
- Full document lifecycle (create, open, save, close, duplicate)
- Reading & searching content (paragraphs, headings, tables, outline)
- Text editing (replace, insert, delete paragraphs and sections)
- Rich formatting (paragraph styles, character runs, alignment, spacing)
- Headers & footers (with page numbers, date fields, images)
- Images (insert from file, resize, alignment)
- Tables (create, read, edit cells, add rows, merge cells, column widths)
- Lists (bullet, numbered, multi-level)
- Document properties (title, author, subject, keywords)
- Page layout (margins, orientation, page size, columns)
- Styles (list, create, modify paragraph/character styles)
- Table of contents generation
- Hyperlinks
- Comments (native Word comments)
- Bookmarks
- Page breaks & section breaks
- Auto-backup before save
- Batch document builder (JSON-driven)
"""

import asyncio
import copy
import json
import os
import re
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Inches, Pt, RGBColor, Emu
from lxml import etree

from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# ---------------------------------------------------------------------------
# Server instance & global state
# ---------------------------------------------------------------------------

server = Server("word-mcp-server")

_current_doc_path: Optional[str] = None
_doc: Optional[Document] = None

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _require_doc() -> Document:
    """Raise an error if no document is currently open."""
    if _doc is None:
        raise RuntimeError("No document open. Use 'open_document' first.")
    return _doc


def _save(backup: bool = True):
    """Save the current document. Optionally create a .bak backup first."""
    doc = _require_doc()
    if backup and _current_doc_path and os.path.exists(_current_doc_path):
        bak_path = _current_doc_path + ".bak"
        shutil.copy2(_current_doc_path, bak_path)
    doc.save(_current_doc_path)


def _para_style(para) -> str:
    return para.style.name if para.style else ""


def _heading_level(para) -> Optional[int]:
    m = re.match(r"Heading (\d+)", _para_style(para))
    return int(m.group(1)) if m else None


def _find_heading_para(doc: Document, heading_text: str):
    """Return the paragraph object matching a heading (case-insensitive partial match)."""
    for p in doc.paragraphs:
        if heading_text.lower() in p.text.lower() and _heading_level(p) is not None:
            return p
    return None


def _find_paras_under_heading(doc: Document, heading_text: str) -> list[int]:
    """
    Return list of paragraph indexes that belong to the section under a heading.
    The section ends when a heading of equal or higher level is encountered.
    """
    paras = doc.paragraphs
    start_idx = None
    heading_lvl = None
    for i, p in enumerate(paras):
        if heading_text.lower() in p.text.lower() and _heading_level(p) is not None:
            start_idx = i
            heading_lvl = _heading_level(p)
            break
    if start_idx is None:
        return []
    result = []
    for i in range(start_idx + 1, len(paras)):
        lvl = _heading_level(paras[i])
        if lvl is not None and lvl <= heading_lvl:
            break
        result.append(i)
    return result


def _delete_paragraph(para):
    """Remove a paragraph element from the document body."""
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _insert_paragraph_after(ref_para, text: str, style_id: Optional[str] = None):
    """Insert a new w:p element immediately after ref_para."""
    new_p = OxmlElement("w:p")
    if style_id:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
        new_p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        new_p.append(r)
    ref_para._p.addnext(new_p)
    return new_p


def _get_style_id(doc: Document, style_name: str) -> Optional[str]:
    """Return the internal style_id for a given style name (case-insensitive)."""
    for s in doc.styles:
        if s.name.lower() == style_name.lower():
            return s.style_id
    return None


def _resolve_style_name(doc: Document, style_name: str) -> Optional[str]:
    """Return the canonical style name if found."""
    for s in doc.styles:
        if s.name.lower() == style_name.lower():
            return s.name
    return None


def _set_run_formatting(run, bold=None, italic=None, underline=None,
                        strike=None, font_size=None, font_name=None,
                        color_hex=None, highlight_color=None,
                        all_caps=None, small_caps=None):
    """Apply character-level formatting to a run."""
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if strike is not None:
        run.font.strike = strike
    if font_size is not None:
        run.font.size = Pt(font_size)
    if font_name is not None:
        run.font.name = font_name
    if color_hex is not None:
        h = color_hex.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    if all_caps is not None:
        run.font.all_caps = all_caps
    if small_caps is not None:
        run.font.small_caps = small_caps


def _xml_to_str(element) -> str:
    return etree.tostring(element, pretty_print=True).decode()


def _make_oxml_element(tag: str) -> Any:
    return OxmlElement(tag)


# ---------------------------------------------------------------------------
# Header / Footer helpers
# ---------------------------------------------------------------------------

def _get_or_create_hdrftr(section, which: str, link_to_previous: bool = False):
    """
    Return the header or footer object for a section.
    which: 'header' | 'footer' | 'first_page_header' | 'first_page_footer'
    """
    attr_map = {
        "header": "header",
        "footer": "footer",
        "first_page_header": "first_page_header",
        "first_page_footer": "first_page_footer",
        "even_page_header": "even_page_header",
        "even_page_footer": "even_page_footer",
    }
    hf = getattr(section, attr_map[which])
    hf.is_linked_to_previous = link_to_previous
    return hf


def _add_page_number_field(paragraph, alignment: str = "CENTER",
                            fmt: str = "PAGE_NUMBER"):
    """
    Insert a { PAGE } or { PAGE } / { NUMPAGES } field into a paragraph.
    fmt: 'PAGE_NUMBER' | 'PAGE_OF_PAGES'
    """
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    paragraph.alignment = align_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.clear()

    if fmt == "PAGE_OF_PAGES":
        # Build "Page X of Y"
        run = paragraph.add_run("Page ")
        _insert_fld_char(paragraph, "PAGE")
        paragraph.add_run(" of ")
        _insert_fld_char(paragraph, "NUMPAGES")
    else:
        _insert_fld_char(paragraph, "PAGE")


def _insert_fld_char(paragraph, instr: str):
    """Insert a simple Word field (PAGE, NUMPAGES, DATE, etc.) into a paragraph."""
    run = OxmlElement("w:r")
    # fldChar begin
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run.append(fld_begin)
    paragraph._p.append(run)

    run2 = OxmlElement("w:r")
    instr_el = OxmlElement("w:instrText")
    instr_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_el.text = f" {instr} "
    run2.append(instr_el)
    paragraph._p.append(run2)

    run3 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3.append(fld_end)
    paragraph._p.append(run3)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str):
    """Set a table cell background shading color (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None,
                     color: str = "000000", size: int = 4):
    """Set individual borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    border_map = {"top": top, "bottom": bottom, "left": left, "right": right}
    for side, val in border_map.items():
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color.lstrip("#"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _merge_cells_horizontal(table, row: int, start_col: int, end_col: int):
    """Merge cells horizontally in a row from start_col to end_col (inclusive)."""
    cell_a = table.cell(row, start_col)
    cell_b = table.cell(row, end_col)
    cell_a.merge(cell_b)


def _merge_cells_vertical(table, col: int, start_row: int, end_row: int):
    """Merge cells vertically in a column from start_row to end_row (inclusive)."""
    cell_a = table.cell(start_row, col)
    cell_b = table.cell(end_row, col)
    cell_a.merge(cell_b)


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, text: str, url: str,
                   color_hex: str = "0563C1", underline: bool = True):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if color_hex:
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color_hex.lstrip("#"))
        rPr.append(color_el)
    if underline:
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rPr.append(u_el)
    run_el.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_el.append(t_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
    return hyperlink


# ---------------------------------------------------------------------------
# Bookmark helpers
# ---------------------------------------------------------------------------

def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int = 1):
    """Wrap the paragraph content in a Word bookmark."""
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, bm_start)
    p.append(bm_end)


# ---------------------------------------------------------------------------
# Comment helpers
# ---------------------------------------------------------------------------

_comment_id_counter = 0


def _add_native_comment(doc: Document, paragraph, text: str,
                         author: str = "Claude", initials: str = "AI"):
    """
    Add a native Word comment to a paragraph (appears in the comments panel).
    This injects the required XML into word/comments.xml part.
    """
    global _comment_id_counter
    _comment_id_counter += 1
    cid = _comment_id_counter

    # Access or create the comments part
    try:
        comments_part = doc.part.comments_part
        comments_el = comments_part._element
    except AttributeError:
        # Create comments part from scratch
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        CT_COMMENTS = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '</w:comments>'
        )
        comments_el = etree.fromstring(CT_COMMENTS.encode())
        # Fallback: insert inline text comment
        paragraph.add_run(f" [COMMENT({author}): {text}]")
        return cid

    date_str = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    comment_xml = (
        f'<w:comment xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:id="{cid}" w:author="{author}" w:date="{date_str}" w:initials="{initials}">'
        f'<w:p><w:r><w:t>{text}</w:t></w:r></w:p>'
        f'</w:comment>'
    )
    comment_el = etree.fromstring(comment_xml.encode())
    comments_el.append(comment_el)

    # Wrap paragraph runs in comment range markers
    p = paragraph._p
    cm_start = OxmlElement("w:commentRangeStart")
    cm_start.set(qn("w:id"), str(cid))
    cm_end = OxmlElement("w:commentRangeEnd")
    cm_end.set(qn("w:id"), str(cid))
    cm_ref_run = OxmlElement("w:r")
    cm_ref = OxmlElement("w:commentReference")
    cm_ref.set(qn("w:id"), str(cid))
    cm_ref_run.append(cm_ref)

    p.insert(0, cm_start)
    p.append(cm_end)
    p.append(cm_ref_run)
    return cid


# ---------------------------------------------------------------------------
# Table of Contents helper
# ---------------------------------------------------------------------------

def _insert_toc(doc: Document, title: str = "Table of Contents",
                max_level: int = 3):
    """
    Insert a Word TOC field at the current end of the document.
    Word will update the TOC on first open (Ctrl+A, F9).
    """
    # Heading for TOC
    if title:
        toc_heading = doc.add_paragraph(title, style="TOC Heading")

    # TOC paragraph with field code
    toc_para = doc.add_paragraph()
    toc_para.style = doc.styles["Normal"]

    run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run.append(fld_begin)
    toc_para._p.append(run)

    run2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
    run2.append(instr)
    toc_para._p.append(run2)

    run3 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3.append(fld_end)
    toc_para._p.append(run3)

    return toc_para


# ---------------------------------------------------------------------------
# Column layout helper
# ---------------------------------------------------------------------------

def _set_section_columns(section, num_cols: int, spacing_cm: float = 1.25,
                          equal_width: bool = True):
    """Set multi-column layout for a document section."""
    sectPr = section._sectPr
    # Remove existing cols element
    existing = sectPr.find(qn("w:cols"))
    if existing is not None:
        sectPr.remove(existing)
    cols_el = OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(num_cols))
    cols_el.set(qn("w:space"), str(int(Cm(spacing_cm).emu / 914)))  # EMU -> twips
    cols_el.set(qn("w:equalWidth"), "1" if equal_width else "0")
    sectPr.append(cols_el)


# ---------------------------------------------------------------------------
# Style creation helper
# ---------------------------------------------------------------------------

def _create_paragraph_style(doc: Document, style_name: str,
                              base_style: str = "Normal",
                              font_name: Optional[str] = None,
                              font_size: Optional[float] = None,
                              bold: Optional[bool] = None,
                              italic: Optional[bool] = None,
                              color_hex: Optional[str] = None,
                              alignment: Optional[str] = None,
                              space_before: Optional[float] = None,
                              space_after: Optional[float] = None):
    """Create a new named paragraph style in the document."""
    # Check if already exists
    existing = _resolve_style_name(doc, style_name)
    if existing:
        style = doc.styles[existing]
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    if base_style:
        try:
            style.base_style = doc.styles[base_style]
        except KeyError:
            pass

    font = style.font
    if font_name:
        font.name = font_name
    if font_size:
        font.size = Pt(font_size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color_hex:
        h = color_hex.lstrip("#")
        font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    pf = style.paragraph_format
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if alignment:
        pf.alignment = align_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

    return style


# ---------------------------------------------------------------------------
# Tool definitions
# ---------------------------------------------------------------------------

@server.list_tools()
async def list_tools() -> list[Tool]:
    return [
        # ---- Document lifecycle ----
        Tool(
            name="open_document",
            description="Open an existing .docx file for editing.",
            inputSchema={
                "type": "object",
                "required": ["path"],
                "properties": {
                    "path": {"type": "string", "description": "Absolute path to the .docx file"}
                }
            }
        ),
        Tool(
            name="create_new_document",
            description="Create a new blank .docx file, optionally from a template.",
            inputSchema={
                "type": "object",
                "required": ["path"],
                "properties": {
                    "path": {"type": "string"},
                    "template": {"type": "string", "description": "Path to a .docx template file"},
                    "default_font": {"type": "string", "description": "Default font name for Normal style"},
                    "default_font_size": {"type": "number", "description": "Default font size in points"}
                }
            }
        ),
        Tool(
            name="save_document",
            description="Save the current document (overwrites original). Creates a .bak backup by default.",
            inputSchema={
                "type": "object",
                "properties": {
                    "backup": {"type": "boolean", "description": "Create .bak backup before saving (default true)"}
                }
            }
        ),
        Tool(
            name="save_as",
            description="Save the current document to a new file path.",
            inputSchema={
                "type": "object",
                "required": ["path"],
                "properties": {"path": {"type": "string"}}
            }
        ),
        Tool(
            name="close_document",
            description="Close the current document without saving.",
            inputSchema={"type": "object", "properties": {}}
        ),
        Tool(
            name="duplicate_document",
            description="Copy the current document file to a new path.",
            inputSchema={
                "type": "object",
                "required": ["new_path"],
                "properties": {"new_path": {"type": "string"}}
            }
        ),

        # ---- Read & inspect ----
        Tool(
            name="get_document_info",
            description="Get document metadata: title, author, dates, word count, paragraph count, sections.",
            inputSchema={"type": "object", "properties": {}}
        ),
        Tool(
            name="read_document",
            description=(
                "Read the full document content as structured JSON. "
                "Returns paragraphs (index, style, text, runs) and tables."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "include_xml": {"type": "boolean", "description": "Include raw XML for each paragraph (default false)"},
                    "include_tables": {"type": "boolean", "description": "Include tables (default true)"}
                }
            }
        ),
        Tool(
            name="get_outline",
            description="Get the document heading hierarchy (outline/TOC preview).",
            inputSchema={"type": "object", "properties": {}}
        ),
        Tool(
            name="get_section",
            description="Get all content paragraphs under a specific heading.",
            inputSchema={
                "type": "object",
                "required": ["heading"],
                "properties": {
                    "heading": {"type": "string", "description": "Heading text (partial match OK)"}
                }
            }
        ),
        Tool(
            name="find_text",
            description="Search for text in the document. Returns all matching paragraphs with index and context.",
            inputSchema={
                "type": "object",
                "required": ["query"],
                "properties": {
                    "query": {"type": "string"},
                    "case_sensitive": {"type": "boolean", "description": "Default false"},
                    "include_tables": {"type": "boolean", "description": "Also search table cells (default true)"}
                }
            }
        ),
        Tool(
            name="read_table",
            description="Read one or all tables in the document.",
            inputSchema={
                "type": "object",
                "properties": {
                    "table_index": {"type": "integer", "description": "Omit to read all tables"}
                }
            }
        ),
        Tool(
            name="list_styles",
            description="List all available styles in the document.",
            inputSchema={
                "type": "object",
                "properties": {
                    "style_type": {
                        "type": "string",
                        "enum": ["paragraph", "character", "table", "all"],
                        "description": "Default: paragraph"
                    }
                }
            }
        ),
        Tool(
            name="get_document_xml",
            description="Get the raw OOXML of a specific paragraph or the entire document body.",
            inputSchema={
                "type": "object",
                "properties": {
                    "paragraph_index": {"type": "integer"},
                    "full_document": {"type": "boolean"}
                }
            }
        ),
        Tool(
            name="read_headers_footers",
            description="Read the content of headers and footers for all sections.",
            inputSchema={"type": "object", "properties": {}}
        ),

        # ---- Text editing ----
        Tool(
            name="replace_text",
            description=(
                "Find and replace text throughout the document. "
                "Preserves run formatting. Searches paragraphs and optionally table cells."
            ),
            inputSchema={
                "type": "object",
                "required": ["find", "replace"],
                "properties": {
                    "find": {"type": "string"},
                    "replace": {"type": "string"},
                    "case_sensitive": {"type": "boolean"},
                    "whole_word": {"type": "boolean"},
                    "include_tables": {"type": "boolean", "description": "Also replace in table cells (default true)"}
                }
            }
        ),
        Tool(
            name="replace_paragraph",
            description="Replace the full text content of a specific paragraph (by index or text match).",
            inputSchema={
                "type": "object",
                "required": ["new_text"],
                "properties": {
                    "match_text": {"type": "string", "description": "Partial text to identify the paragraph"},
                    "paragraph_index": {"type": "integer"},
                    "new_text": {"type": "string"},
                    "preserve_style": {"type": "boolean", "description": "Keep original style (default true)"}
                }
            }
        ),
        Tool(
            name="replace_section",
            description="Replace all content paragraphs under a heading with new content.",
            inputSchema={
                "type": "object",
                "required": ["heading", "new_paragraphs"],
                "properties": {
                    "heading": {"type": "string"},
                    "new_paragraphs": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "text": {"type": "string"},
                                "style": {"type": "string"}
                            }
                        }
                    }
                }
            }
        ),
        Tool(
            name="insert_paragraph",
            description=(
                "Insert a new paragraph at a given position. "
                "Position can be: after_text, after_heading, at_index, or at_end."
            ),
            inputSchema={
                "type": "object",
                "required": ["text"],
                "properties": {
                    "text": {"type": "string"},
                    "style": {"type": "string", "description": "Style name e.g. 'Normal', 'Heading 1', 'List Bullet'"},
                    "after_text": {"type": "string"},
                    "after_heading": {"type": "string"},
                    "at_index": {"type": "integer"},
                    "at_end": {"type": "boolean"},
                    "bold": {"type": "boolean"},
                    "italic": {"type": "boolean"},
                    "font_size": {"type": "number"},
                    "font_name": {"type": "string"},
                    "color_hex": {"type": "string"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT", "JUSTIFY"]}
                }
            }
        ),
        Tool(
            name="delete_paragraph",
            description="Delete one or more paragraphs by index or text match.",
            inputSchema={
                "type": "object",
                "properties": {
                    "match_text": {"type": "string"},
                    "paragraph_index": {"type": "integer"},
                    "delete_all_matching": {"type": "boolean", "description": "Delete all matches (default false, only first)"}
                }
            }
        ),
        Tool(
            name="delete_section",
            description="Delete a heading paragraph and all content under it.",
            inputSchema={
                "type": "object",
                "required": ["heading"],
                "properties": {"heading": {"type": "string"}}
            }
        ),
        Tool(
            name="move_section",
            description="Move a section (heading + its content) to before or after another heading.",
            inputSchema={
                "type": "object",
                "required": ["section_heading"],
                "properties": {
                    "section_heading": {"type": "string"},
                    "before_heading": {"type": "string"},
                    "after_heading": {"type": "string"}
                }
            }
        ),

        # ---- Formatting ----
        Tool(
            name="format_paragraph",
            description="Change paragraph-level formatting: style, alignment, spacing, indentation.",
            inputSchema={
                "type": "object",
                "properties": {
                    "match_text": {"type": "string"},
                    "paragraph_index": {"type": "integer"},
                    "style": {"type": "string"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT", "JUSTIFY"]},
                    "space_before": {"type": "number", "description": "Points before paragraph"},
                    "space_after": {"type": "number", "description": "Points after paragraph"},
                    "left_indent": {"type": "number", "description": "Left indent in cm"},
                    "right_indent": {"type": "number", "description": "Right indent in cm"},
                    "first_line_indent": {"type": "number", "description": "First-line indent in cm"},
                    "line_spacing": {"type": "number", "description": "Line spacing in points (e.g. 24 = double)"},
                    "keep_together": {"type": "boolean"},
                    "keep_with_next": {"type": "boolean"},
                    "page_break_before": {"type": "boolean"}
                }
            }
        ),
        Tool(
            name="format_text_run",
            description=(
                "Apply character-level formatting to runs within a paragraph. "
                "Can target a specific run by text match."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "paragraph_match": {"type": "string", "description": "Text to find the paragraph"},
                    "paragraph_index": {"type": "integer"},
                    "run_text_match": {"type": "string", "description": "Narrow to runs containing this text"},
                    "bold": {"type": "boolean"},
                    "italic": {"type": "boolean"},
                    "underline": {"type": "boolean"},
                    "strike": {"type": "boolean"},
                    "font_size": {"type": "number"},
                    "font_name": {"type": "string"},
                    "color_hex": {"type": "string", "description": "6-char hex e.g. 'FF0000'"},
                    "all_caps": {"type": "boolean"},
                    "small_caps": {"type": "boolean"}
                }
            }
        ),

        # ---- Headings ----
        Tool(
            name="add_heading",
            description="Add a new heading at the end of the document or after a specific heading.",
            inputSchema={
                "type": "object",
                "required": ["text", "level"],
                "properties": {
                    "text": {"type": "string"},
                    "level": {"type": "integer", "description": "Heading level 1-9"},
                    "after_heading": {"type": "string", "description": "Insert after this heading"},
                    "at_end": {"type": "boolean", "description": "Append at end (default true)"}
                }
            }
        ),

        # ---- Tables ----
        Tool(
            name="insert_table",
            description=(
                "Insert a table with optional data, style, and column widths. "
                "Supports header row formatting and cell-level formatting via dict values."
            ),
            inputSchema={
                "type": "object",
                "required": ["rows", "cols"],
                "properties": {
                    "rows": {"type": "integer"},
                    "cols": {"type": "integer"},
                    "data": {
                        "type": "array",
                        "description": "Array of rows; each row is an array of strings or {text, bold, color_hex} dicts",
                        "items": {"type": "array"}
                    },
                    "style": {"type": "string", "description": "Table style name (default: 'Table Grid')"},
                    "header_row": {"type": "boolean", "description": "Make first row bold"},
                    "col_widths": {"type": "array", "items": {"type": "number"}, "description": "Column widths in cm"},
                    "after_text": {"type": "string", "description": "Insert after paragraph containing this text"},
                    "after_heading": {"type": "string"}
                }
            }
        ),
        Tool(
            name="edit_table_cell",
            description="Edit the text of a specific table cell.",
            inputSchema={
                "type": "object",
                "required": ["table_index", "row", "col", "text"],
                "properties": {
                    "table_index": {"type": "integer"},
                    "row": {"type": "integer"},
                    "col": {"type": "integer"},
                    "text": {"type": "string"},
                    "bold": {"type": "boolean"},
                    "italic": {"type": "boolean"},
                    "font_size": {"type": "number"},
                    "color_hex": {"type": "string"},
                    "bg_color": {"type": "string", "description": "Cell background color hex (no #)"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT", "JUSTIFY"]}
                }
            }
        ),
        Tool(
            name="add_table_row",
            description="Append a new row to an existing table.",
            inputSchema={
                "type": "object",
                "required": ["table_index"],
                "properties": {
                    "table_index": {"type": "integer"},
                    "data": {"type": "array", "items": {"type": "string"}}
                }
            }
        ),
        Tool(
            name="merge_table_cells",
            description="Merge table cells horizontally or vertically.",
            inputSchema={
                "type": "object",
                "required": ["table_index", "direction"],
                "properties": {
                    "table_index": {"type": "integer"},
                    "direction": {"type": "string", "enum": ["horizontal", "vertical"]},
                    "row": {"type": "integer", "description": "Row index (for horizontal merge)"},
                    "col": {"type": "integer", "description": "Column index (for vertical merge)"},
                    "start_col": {"type": "integer"},
                    "end_col": {"type": "integer"},
                    "start_row": {"type": "integer"},
                    "end_row": {"type": "integer"}
                }
            }
        ),
        Tool(
            name="format_table_cell",
            description="Apply background color and/or borders to a table cell.",
            inputSchema={
                "type": "object",
                "required": ["table_index", "row", "col"],
                "properties": {
                    "table_index": {"type": "integer"},
                    "row": {"type": "integer"},
                    "col": {"type": "integer"},
                    "bg_color": {"type": "string", "description": "Fill color hex e.g. 'FFCC00'"},
                    "border_top": {"type": "string", "enum": ["single", "double", "none"]},
                    "border_bottom": {"type": "string", "enum": ["single", "double", "none"]},
                    "border_left": {"type": "string", "enum": ["single", "double", "none"]},
                    "border_right": {"type": "string", "enum": ["single", "double", "none"]},
                    "border_color": {"type": "string", "description": "Border color hex (default 000000)"},
                    "border_size": {"type": "integer", "description": "Border size in eighths of a point (default 4)"}
                }
            }
        ),
        Tool(
            name="delete_table_row",
            description="Delete a row from a table.",
            inputSchema={
                "type": "object",
                "required": ["table_index", "row"],
                "properties": {
                    "table_index": {"type": "integer"},
                    "row": {"type": "integer"}
                }
            }
        ),
        Tool(
            name="delete_table",
            description="Delete an entire table from the document.",
            inputSchema={
                "type": "object",
                "required": ["table_index"],
                "properties": {
                    "table_index": {"type": "integer"}
                }
            }
        ),

        # ---- Headers & Footers ----
        Tool(
            name="set_header",
            description=(
                "Set the header text for a document section. "
                "Supports left/center/right tabs, page numbers, and date fields."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "Header text"},
                    "section_index": {"type": "integer", "description": "Section index (default 0)"},
                    "which": {
                        "type": "string",
                        "enum": ["header", "first_page_header", "even_page_header"],
                        "description": "Default: header"
                    },
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT"]},
                    "bold": {"type": "boolean"},
                    "italic": {"type": "boolean"},
                    "font_size": {"type": "number"},
                    "font_name": {"type": "string"}
                }
            }
        ),
        Tool(
            name="set_footer",
            description=(
                "Set the footer text for a document section. "
                "Supports page numbers and custom text."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "Footer text (use '' to insert only page number)"},
                    "section_index": {"type": "integer", "description": "Section index (default 0)"},
                    "which": {
                        "type": "string",
                        "enum": ["footer", "first_page_footer", "even_page_footer"],
                        "description": "Default: footer"
                    },
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT"]},
                    "add_page_number": {"type": "boolean", "description": "Add PAGE field"},
                    "page_number_format": {
                        "type": "string",
                        "enum": ["PAGE_NUMBER", "PAGE_OF_PAGES"],
                        "description": "Page number format"
                    },
                    "bold": {"type": "boolean"},
                    "font_size": {"type": "number"},
                    "font_name": {"type": "string"}
                }
            }
        ),
        Tool(
            name="add_image_to_header",
            description="Insert an image (e.g. company logo) into a header.",
            inputSchema={
                "type": "object",
                "required": ["image_path"],
                "properties": {
                    "image_path": {"type": "string", "description": "Absolute path to image file (PNG, JPG, etc.)"},
                    "section_index": {"type": "integer", "description": "Default 0"},
                    "width_cm": {"type": "number", "description": "Image width in cm"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT"]}
                }
            }
        ),
        Tool(
            name="clear_header",
            description="Remove all content from a header.",
            inputSchema={
                "type": "object",
                "properties": {
                    "section_index": {"type": "integer", "description": "Default 0"},
                    "which": {"type": "string", "enum": ["header", "first_page_header", "even_page_header"]}
                }
            }
        ),
        Tool(
            name="clear_footer",
            description="Remove all content from a footer.",
            inputSchema={
                "type": "object",
                "properties": {
                    "section_index": {"type": "integer", "description": "Default 0"},
                    "which": {"type": "string", "enum": ["footer", "first_page_footer", "even_page_footer"]}
                }
            }
        ),

        # ---- Images ----
        Tool(
            name="insert_image",
            description="Insert an image into the document body at a given position.",
            inputSchema={
                "type": "object",
                "required": ["image_path"],
                "properties": {
                    "image_path": {"type": "string", "description": "Absolute path to image file"},
                    "width_cm": {"type": "number", "description": "Width in cm (height auto-scaled)"},
                    "height_cm": {"type": "number", "description": "Height in cm (optional)"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT"]},
                    "after_text": {"type": "string"},
                    "after_heading": {"type": "string"},
                    "at_end": {"type": "boolean"},
                    "caption": {"type": "string", "description": "Optional caption paragraph below image"}
                }
            }
        ),

        # ---- Page layout ----
        Tool(
            name="set_page_margins",
            description="Set page margins for all or a specific section.",
            inputSchema={
                "type": "object",
                "properties": {
                    "top": {"type": "number", "description": "Top margin in cm"},
                    "bottom": {"type": "number", "description": "Bottom margin in cm"},
                    "left": {"type": "number", "description": "Left margin in cm"},
                    "right": {"type": "number", "description": "Right margin in cm"},
                    "section_index": {"type": "integer", "description": "Omit to apply to all sections"}
                }
            }
        ),
        Tool(
            name="set_page_orientation",
            description="Set page orientation for a section.",
            inputSchema={
                "type": "object",
                "required": ["orientation"],
                "properties": {
                    "orientation": {"type": "string", "enum": ["portrait", "landscape"]},
                    "section_index": {"type": "integer", "description": "Default 0"}
                }
            }
        ),
        Tool(
            name="set_page_size",
            description="Set the page size (e.g. A4, Letter, or custom dimensions).",
            inputSchema={
                "type": "object",
                "properties": {
                    "preset": {
                        "type": "string",
                        "enum": ["A4", "A3", "A5", "Letter", "Legal"],
                        "description": "Use a preset size"
                    },
                    "width_cm": {"type": "number", "description": "Custom width in cm"},
                    "height_cm": {"type": "number", "description": "Custom height in cm"},
                    "section_index": {"type": "integer", "description": "Default 0"}
                }
            }
        ),
        Tool(
            name="add_section_break",
            description="Insert a section break (new page, continuous, even page, odd page).",
            inputSchema={
                "type": "object",
                "properties": {
                    "break_type": {
                        "type": "string",
                        "enum": ["new_page", "continuous", "even_page", "odd_page"],
                        "description": "Default: new_page"
                    },
                    "after_text": {"type": "string"},
                    "at_end": {"type": "boolean"}
                }
            }
        ),
        Tool(
            name="add_page_break",
            description="Insert a hard page break after a specific paragraph.",
            inputSchema={
                "type": "object",
                "properties": {
                    "after_text": {"type": "string"},
                    "at_index": {"type": "integer"},
                    "at_end": {"type": "boolean"}
                }
            }
        ),
        Tool(
            name="set_columns",
            description="Set multi-column layout for a section.",
            inputSchema={
                "type": "object",
                "required": ["num_cols"],
                "properties": {
                    "num_cols": {"type": "integer", "description": "Number of columns (1 to disable)"},
                    "spacing_cm": {"type": "number", "description": "Space between columns in cm (default 1.25)"},
                    "section_index": {"type": "integer", "description": "Default 0"}
                }
            }
        ),

        # ---- Styles ----
        Tool(
            name="create_style",
            description="Create or update a custom paragraph style.",
            inputSchema={
                "type": "object",
                "required": ["style_name"],
                "properties": {
                    "style_name": {"type": "string"},
                    "base_style": {"type": "string", "description": "Parent style (default: Normal)"},
                    "font_name": {"type": "string"},
                    "font_size": {"type": "number"},
                    "bold": {"type": "boolean"},
                    "italic": {"type": "boolean"},
                    "color_hex": {"type": "string"},
                    "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT", "JUSTIFY"]},
                    "space_before": {"type": "number", "description": "Points"},
                    "space_after": {"type": "number", "description": "Points"}
                }
            }
        ),

        # ---- TOC ----
        Tool(
            name="insert_toc",
            description=(
                "Insert a Table of Contents field. "
                "Word will render it on first open (press Ctrl+A then F9 to update)."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "TOC heading text (default: 'Table of Contents')"},
                    "max_level": {"type": "integer", "description": "Max heading level to include (default 3)"},
                    "after_heading": {"type": "string"},
                    "at_end": {"type": "boolean"}
                }
            }
        ),

        # ---- Hyperlinks ----
        Tool(
            name="add_hyperlink",
            description="Add a hyperlink to an existing paragraph.",
            inputSchema={
                "type": "object",
                "required": ["paragraph_match", "text", "url"],
                "properties": {
                    "paragraph_match": {"type": "string", "description": "Text to identify the target paragraph"},
                    "paragraph_index": {"type": "integer"},
                    "text": {"type": "string", "description": "Display text for the hyperlink"},
                    "url": {"type": "string", "description": "URL for the hyperlink"},
                    "color_hex": {"type": "string", "description": "Link color hex (default: 0563C1)"},
                    "underline": {"type": "boolean", "description": "Default true"}
                }
            }
        ),

        # ---- Comments ----
        Tool(
            name="add_comment",
            description="Add a native Word comment to a paragraph (visible in the review pane).",
            inputSchema={
                "type": "object",
                "required": ["match_text", "comment"],
                "properties": {
                    "match_text": {"type": "string"},
                    "comment": {"type": "string"},
                    "author": {"type": "string", "description": "Comment author (default: Claude)"},
                    "initials": {"type": "string", "description": "Author initials (default: AI)"}
                }
            }
        ),

        # ---- Bookmarks ----
        Tool(
            name="add_bookmark",
            description="Add a named bookmark to a paragraph.",
            inputSchema={
                "type": "object",
                "required": ["match_text", "bookmark_name"],
                "properties": {
                    "match_text": {"type": "string"},
                    "paragraph_index": {"type": "integer"},
                    "bookmark_name": {"type": "string"},
                    "bookmark_id": {"type": "integer", "description": "Unique numeric ID (default: 1)"}
                }
            }
        ),

        # ---- Document properties ----
        Tool(
            name="set_document_properties",
            description="Set document core properties (title, author, subject, keywords, description, category).",
            inputSchema={
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "author": {"type": "string"},
                    "subject": {"type": "string"},
                    "keywords": {"type": "string"},
                    "description": {"type": "string"},
                    "category": {"type": "string"},
                    "company": {"type": "string"}
                }
            }
        ),

        # ---- Advanced XML ----
        Tool(
            name="apply_xml_patch",
            description="Replace a paragraph's XML with custom OOXML (for advanced formatting not covered by other tools).",
            inputSchema={
                "type": "object",
                "required": ["paragraph_index", "xml_content"],
                "properties": {
                    "paragraph_index": {"type": "integer"},
                    "xml_content": {"type": "string", "description": "Full XML of the new w:p element"}
                }
            }
        ),

        # ---- Batch builder ----
        Tool(
            name="build_document",
            description=(
                "Create a fully-formatted document in ONE call from a structured JSON spec. "
                "Supports: headings (H1-H9), paragraphs, bullet/numbered lists, tables, "
                "page breaks, section breaks, images, hyperlinks, per-run formatting "
                "(bold, italic, underline, strikethrough, font, size, color), "
                "header/footer, TOC, and document properties. "
                "Use this for creating new documents with substantial content."
            ),
            inputSchema={
                "type": "object",
                "required": ["path", "elements"],
                "properties": {
                    "path": {"type": "string"},
                    "template": {"type": "string", "description": "Path to .docx template"},
                    "default_font": {"type": "string"},
                    "default_font_size": {"type": "number"},
                    "title": {"type": "string"},
                    "author": {"type": "string"},
                    "subject": {"type": "string"},
                    "header_text": {"type": "string", "description": "Global header for all sections"},
                    "footer_text": {"type": "string"},
                    "footer_page_numbers": {"type": "boolean", "description": "Add page numbers to footer"},
                    "margins": {
                        "type": "object",
                        "properties": {
                            "top": {"type": "number"}, "bottom": {"type": "number"},
                            "left": {"type": "number"}, "right": {"type": "number"}
                        }
                    },
                    "elements": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "type": {
                                    "type": "string",
                                    "enum": ["heading", "paragraph", "list", "table",
                                             "page_break", "section_break", "image",
                                             "toc", "hyperlink"]
                                },
                                "text": {"type": "string"},
                                "level": {"type": "integer"},
                                "style": {"type": "string"},
                                "alignment": {"type": "string", "enum": ["LEFT", "CENTER", "RIGHT", "JUSTIFY"]},
                                "space_before": {"type": "number"},
                                "space_after": {"type": "number"},
                                "runs": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "text": {"type": "string"},
                                            "bold": {"type": "boolean"},
                                            "italic": {"type": "boolean"},
                                            "underline": {"type": "boolean"},
                                            "strike": {"type": "boolean"},
                                            "font_size": {"type": "number"},
                                            "font_name": {"type": "string"},
                                            "color_hex": {"type": "string"},
                                            "url": {"type": "string", "description": "Makes this run a hyperlink"}
                                        }
                                    }
                                },
                                "list_type": {"type": "string", "enum": ["bullet", "numbered"]},
                                "list_level": {"type": "integer"},
                                "data": {"type": "array"},
                                "header_row": {"type": "boolean"},
                                "table_style": {"type": "string"},
                                "col_widths": {"type": "array"},
                                "image_path": {"type": "string"},
                                "width_cm": {"type": "number"},
                                "caption": {"type": "string"},
                                "url": {"type": "string"},
                                "toc_title": {"type": "string"},
                                "toc_max_level": {"type": "integer"},
                                "break_type": {"type": "string"}
                            }
                        }
                    }
                }
            }
        ),
    ]


# ---------------------------------------------------------------------------
# Tool implementations
# ---------------------------------------------------------------------------

@server.call_tool()
async def call_tool(name: str, arguments: dict[str, Any]) -> list[TextContent]:
    global _current_doc_path, _doc

    def ok(msg):
        return [TextContent(type="text", text=str(msg))]

    # Shared alignment map
    ALIGN_MAP = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    try:

        # ====================================================================
        # Document lifecycle
        # ====================================================================

        if name == "open_document":
            path = arguments["path"]
            if not os.path.exists(path):
                return ok(f"ERROR: File not found: {path}")
            _current_doc_path = path
            _doc = Document(path)
            return ok(
                f"Opened: {path}\n"
                f"Paragraphs: {len(_doc.paragraphs)} | "
                f"Tables: {len(_doc.tables)} | "
                f"Sections: {len(_doc.sections)}"
            )

        elif name == "create_new_document":
            path = arguments["path"]
            template = arguments.get("template")
            doc = Document(template) if template else Document()
            # Apply default font/size if specified
            if arguments.get("default_font") or arguments.get("default_font_size"):
                style = doc.styles["Normal"]
                if arguments.get("default_font"):
                    style.font.name = arguments["default_font"]
                if arguments.get("default_font_size"):
                    style.font.size = Pt(arguments["default_font_size"])
            doc.save(path)
            _current_doc_path = path
            _doc = Document(path)
            return ok(f"Created: {path}")

        elif name == "save_document":
            backup = arguments.get("backup", True)
            _save(backup=backup)
            return ok(f"Saved: {_current_doc_path}")

        elif name == "save_as":
            path = arguments["path"]
            _require_doc().save(path)
            _current_doc_path = path
            return ok(f"Saved as: {path}")

        elif name == "close_document":
            _doc = None
            _current_doc_path = None
            return ok("Document closed.")

        elif name == "duplicate_document":
            doc = _require_doc()
            new_path = arguments["new_path"]
            shutil.copy2(_current_doc_path, new_path)
            return ok(f"Duplicated to: {new_path}")

        # ====================================================================
        # Read & inspect
        # ====================================================================

        elif name == "get_document_info":
            doc = _require_doc()
            cp = doc.core_properties
            words = sum(len(p.text.split()) for p in doc.paragraphs)
            chars = sum(len(p.text) for p in doc.paragraphs)
            result = {
                "path": _current_doc_path,
                "title": cp.title,
                "author": cp.author,
                "last_modified_by": cp.last_modified_by,
                "created": str(cp.created),
                "modified": str(cp.modified),
                "subject": cp.subject,
                "keywords": cp.keywords,
                "description": cp.description,
                "category": cp.category,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "word_count": words,
                "char_count": chars,
            }
            return ok(json.dumps(result, indent=2, default=str))

        elif name == "read_document":
            doc = _require_doc()
            include_xml = arguments.get("include_xml", False)
            include_tables = arguments.get("include_tables", True)
            elements = []
            for i, p in enumerate(doc.paragraphs):
                el = {
                    "index": i,
                    "type": "paragraph",
                    "style": _para_style(p),
                    "text": p.text,
                    "alignment": str(p.alignment),
                    "runs": [
                        {
                            "text": r.text,
                            "bold": r.bold,
                            "italic": r.italic,
                            "underline": r.underline,
                            "font_size": r.font.size.pt if r.font.size else None,
                            "font_name": r.font.name,
                            "color": str(r.font.color.rgb) if r.font.color and r.font.color.type else None,
                        }
                        for r in p.runs
                    ],
                }
                if include_xml:
                    el["xml"] = _xml_to_str(p._p)
                elements.append(el)
            tables = []
            if include_tables:
                for ti, tbl in enumerate(doc.tables):
                    rows_data = []
                    for row in tbl.rows:
                        rows_data.append([cell.text for cell in row.cells])
                    tables.append({
                        "table_index": ti,
                        "style": tbl.style.name if tbl.style else None,
                        "rows": rows_data
                    })
            return ok(json.dumps({"paragraphs": elements, "tables": tables}, indent=2, default=str))

        elif name == "get_outline":
            doc = _require_doc()
            outline = []
            for i, p in enumerate(doc.paragraphs):
                lvl = _heading_level(p)
                if lvl is not None:
                    outline.append({
                        "index": i, "level": lvl, "text": p.text
                    })
            lines = [
                f"[{o['index']}] {'  ' * (o['level'] - 1)}H{o['level']}: {o['text']}"
                for o in outline
            ]
            return ok("\n".join(lines) if lines else "No headings found.")

        elif name == "get_section":
            doc = _require_doc()
            heading = arguments["heading"]
            idxs = _find_paras_under_heading(doc, heading)
            if not idxs:
                return ok(f"Heading '{heading}' not found or section is empty.")
            paras = doc.paragraphs
            lines = [f"[{i}] ({_para_style(paras[i])}) {paras[i].text}" for i in idxs]
            return ok("\n".join(lines))

        elif name == "find_text":
            doc = _require_doc()
            query = arguments["query"]
            case_sensitive = arguments.get("case_sensitive", False)
            include_tables = arguments.get("include_tables", True)
            results = []
            for i, p in enumerate(doc.paragraphs):
                haystack = p.text if case_sensitive else p.text.lower()
                needle = query if case_sensitive else query.lower()
                if needle in haystack:
                    results.append({"source": "paragraph", "index": i,
                                    "style": _para_style(p), "text": p.text})
            if include_tables:
                for ti, tbl in enumerate(doc.tables):
                    for ri, row in enumerate(tbl.rows):
                        for ci, cell in enumerate(row.cells):
                            haystack = cell.text if case_sensitive else cell.text.lower()
                            needle = query if case_sensitive else query.lower()
                            if needle in haystack:
                                results.append({
                                    "source": "table",
                                    "table_index": ti,
                                    "row": ri, "col": ci,
                                    "text": cell.text
                                })
            if not results:
                return ok(f"No occurrences of '{query}' found.")
            return ok(json.dumps(results, indent=2))

        elif name == "read_table":
            doc = _require_doc()
            ti = arguments.get("table_index")
            tables = doc.tables
            if ti is not None:
                selected = [tables[ti]]
                base_idx = ti
            else:
                selected = tables
                base_idx = 0
            result = []
            for idx, tbl in enumerate(selected):
                rows_data = [[cell.text for cell in row.cells] for row in tbl.rows]
                result.append({
                    "index": idx + base_idx,
                    "style": tbl.style.name if tbl.style else None,
                    "rows": len(tbl.rows),
                    "cols": len(tbl.columns),
                    "data": rows_data
                })
            return ok(json.dumps(result, indent=2))

        elif name == "list_styles":
            doc = _require_doc()
            stype = arguments.get("style_type", "paragraph")
            type_map = {
                "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                "character": WD_STYLE_TYPE.CHARACTER,
                "table": WD_STYLE_TYPE.TABLE,
            }
            if stype == "all":
                styles = [{"name": s.name, "type": str(s.type), "id": s.style_id}
                          for s in doc.styles]
            else:
                styles = [{"name": s.name, "id": s.style_id}
                          for s in doc.styles if s.type == type_map.get(stype)]
            return ok(json.dumps(styles, indent=2))

        elif name == "get_document_xml":
            doc = _require_doc()
            if arguments.get("full_document"):
                return ok(_xml_to_str(doc.element.body))
            idx = arguments.get("paragraph_index")
            if idx is not None:
                return ok(_xml_to_str(doc.paragraphs[idx]._p))
            return ok("Specify paragraph_index or full_document=true")

        elif name == "read_headers_footers":
            doc = _require_doc()
            result = []
            for si, section in enumerate(doc.sections):
                entry = {"section": si}
                for which in ["header", "footer", "first_page_header", "first_page_footer",
                              "even_page_header", "even_page_footer"]:
                    try:
                        hf = getattr(section, which)
                        if hf and not hf.is_linked_to_previous:
                            text = "\n".join(p.text for p in hf.paragraphs)
                            entry[which] = text
                    except Exception:
                        pass
                result.append(entry)
            return ok(json.dumps(result, indent=2))

        # ====================================================================
        # Text editing
        # ====================================================================

        elif name == "replace_text":
            doc = _require_doc()
            find = arguments["find"]
            replace_with = arguments["replace"]
            case_sensitive = arguments.get("case_sensitive", False)
            whole_word = arguments.get("whole_word", False)
            include_tables = arguments.get("include_tables", True)
            count = 0

            def do_replace(runs):
                nonlocal count
                for run in runs:
                    orig = run.text
                    text_cmp = orig if case_sensitive else orig.lower()
                    needle = find if case_sensitive else find.lower()
                    if whole_word:
                        pattern = rf"\b{re.escape(needle)}\b"
                    else:
                        pattern = re.escape(needle)
                    flags = 0 if case_sensitive else re.IGNORECASE
                    if re.search(pattern, text_cmp):
                        run.text = re.sub(
                            rf"\b{re.escape(find)}\b" if whole_word else re.escape(find),
                            replace_with, orig, flags=flags
                        )
                        count += 1

            for para in doc.paragraphs:
                do_replace(para.runs)
            if include_tables:
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                do_replace(para.runs)
            _save()
            return ok(f"Replaced {count} occurrence(s) of '{find}' -> '{replace_with}'")

        elif name == "replace_paragraph":
            doc = _require_doc()
            new_text = arguments["new_text"]
            target = None
            if "paragraph_index" in arguments:
                target = doc.paragraphs[arguments["paragraph_index"]]
            elif "match_text" in arguments:
                for p in doc.paragraphs:
                    if arguments["match_text"].lower() in p.text.lower():
                        target = p
                        break
            if target is None:
                return ok("Paragraph not found.")
            for run in target.runs:
                run.text = ""
            if target.runs:
                target.runs[0].text = new_text
            else:
                target.add_run(new_text)
            _save()
            return ok("Paragraph replaced.")

        elif name == "replace_section":
            doc = _require_doc()
            heading = arguments["heading"]
            new_paragraphs = arguments["new_paragraphs"]
            idxs = _find_paras_under_heading(doc, heading)
            if not idxs:
                return ok(f"Heading '{heading}' not found.")
            paras = doc.paragraphs
            for i in reversed(idxs):
                _delete_paragraph(paras[i])
            # Re-find heading after deletions
            heading_para = _find_heading_para(doc, heading)
            if heading_para is None:
                return ok("Heading paragraph no longer found after deletion.")
            ref = heading_para
            for np_data in new_paragraphs:
                style_name = np_data.get("style", "Normal")
                style_id = _get_style_id(doc, style_name) or "Normal"
                _insert_paragraph_after(ref, np_data["text"], style_id)
                # Advance ref to the newly inserted paragraph
                idx_ref = list(doc.paragraphs).index(ref)
                if idx_ref + 1 < len(doc.paragraphs):
                    ref = doc.paragraphs[idx_ref + 1]
            _save()
            return ok(f"Section '{heading}' replaced with {len(new_paragraphs)} paragraph(s).")

        elif name == "insert_paragraph":
            doc = _require_doc()
            text = arguments["text"]
            style = arguments.get("style", "Normal")
            style_id = _get_style_id(doc, style) or "Normal"
            paras = doc.paragraphs

            ref_para = None
            if "after_text" in arguments:
                for p in paras:
                    if arguments["after_text"].lower() in p.text.lower():
                        ref_para = p
                        break
            elif "after_heading" in arguments:
                ref_para = _find_heading_para(doc, arguments["after_heading"])
            elif "at_index" in arguments:
                idx = arguments["at_index"]
                if 0 < idx < len(paras):
                    ref_para = paras[idx - 1]

            if ref_para is not None:
                new_p = _insert_paragraph_after(ref_para, text, style_id)
                # Find the newly created paragraph object
                new_para = None
                for p in doc.paragraphs:
                    if p._p is new_p:
                        new_para = p
                        break
            else:
                # Append at end
                new_para = doc.add_paragraph(text)
                try:
                    new_para.style = doc.styles[style]
                except Exception:
                    pass

            # Apply optional run-level formatting if paragraph found
            if new_para:
                align_str = arguments.get("alignment")
                if align_str:
                    new_para.alignment = ALIGN_MAP.get(align_str.upper(), WD_ALIGN_PARAGRAPH.LEFT)
                for run in new_para.runs:
                    _set_run_formatting(
                        run,
                        bold=arguments.get("bold"),
                        italic=arguments.get("italic"),
                        font_size=arguments.get("font_size"),
                        font_name=arguments.get("font_name"),
                        color_hex=arguments.get("color_hex"),
                    )
            _save()
            return ok("Paragraph inserted.")

        elif name == "delete_paragraph":
            doc = _require_doc()
            paras = doc.paragraphs
            if "paragraph_index" in arguments:
                _delete_paragraph(paras[arguments["paragraph_index"]])
                _save()
                return ok("Paragraph deleted.")
            if "match_text" in arguments:
                delete_all = arguments.get("delete_all_matching", False)
                to_delete = []
                for p in paras:
                    if arguments["match_text"].lower() in p.text.lower():
                        to_delete.append(p)
                        if not delete_all:
                            break
                for p in to_delete:
                    _delete_paragraph(p)
                _save()
                return ok(f"Deleted {len(to_delete)} paragraph(s).")
            return ok("Specify match_text or paragraph_index.")

        elif name == "delete_section":
            doc = _require_doc()
            heading = arguments["heading"]
            heading_para = _find_heading_para(doc, heading)
            if heading_para is None:
                return ok(f"Heading '{heading}' not found.")
            idxs = _find_paras_under_heading(doc, heading)
            for i in reversed(idxs):
                _delete_paragraph(doc.paragraphs[i])
            _delete_paragraph(heading_para)
            _save()
            return ok(f"Section '{heading}' deleted ({len(idxs) + 1} paragraphs removed).")

        elif name == "move_section":
            doc = _require_doc()
            section_heading = arguments["section_heading"]
            before_heading = arguments.get("before_heading")
            after_heading_arg = arguments.get("after_heading")

            heading_para = _find_heading_para(doc, section_heading)
            if heading_para is None:
                return ok(f"Section heading '{section_heading}' not found.")
            idxs = _find_paras_under_heading(doc, section_heading)
            all_paras = [heading_para] + [doc.paragraphs[i] for i in idxs]
            clones = [copy.deepcopy(p._p) for p in all_paras]
            for p in reversed(all_paras):
                _delete_paragraph(p)

            target_text = before_heading or after_heading_arg
            target_para = _find_heading_para(doc, target_text)
            if target_para is None:
                return ok(f"Target heading '{target_text}' not found.")

            if before_heading:
                ref = target_para._p
                for clone in reversed(clones):
                    ref.addprevious(clone)
            else:
                ref = target_para._p
                for clone in clones:
                    ref.addnext(clone)
                    ref = clone
            _save()
            return ok(f"Section '{section_heading}' moved.")

        # ====================================================================
        # Formatting
        # ====================================================================

        elif name == "format_paragraph":
            doc = _require_doc()
            target = None
            if "paragraph_index" in arguments:
                target = doc.paragraphs[arguments["paragraph_index"]]
            elif "match_text" in arguments:
                for p in doc.paragraphs:
                    if arguments["match_text"].lower() in p.text.lower():
                        target = p
                        break
            if target is None:
                return ok("Paragraph not found.")

            if "style" in arguments:
                try:
                    target.style = doc.styles[arguments["style"]]
                except KeyError:
                    pass
            if "alignment" in arguments:
                target.alignment = ALIGN_MAP.get(arguments["alignment"].upper(), WD_ALIGN_PARAGRAPH.LEFT)

            pf = target.paragraph_format
            if "space_before" in arguments:
                pf.space_before = Pt(arguments["space_before"])
            if "space_after" in arguments:
                pf.space_after = Pt(arguments["space_after"])
            if "left_indent" in arguments:
                pf.left_indent = Cm(arguments["left_indent"])
            if "right_indent" in arguments:
                pf.right_indent = Cm(arguments["right_indent"])
            if "first_line_indent" in arguments:
                pf.first_line_indent = Cm(arguments["first_line_indent"])
            if "line_spacing" in arguments:
                pf.line_spacing = Pt(arguments["line_spacing"])
            if "keep_together" in arguments:
                pf.keep_together = arguments["keep_together"]
            if "keep_with_next" in arguments:
                pf.keep_with_next = arguments["keep_with_next"]
            if "page_break_before" in arguments:
                pf.page_break_before = arguments["page_break_before"]

            _save()
            return ok("Paragraph formatted.")

        elif name == "format_text_run":
            doc = _require_doc()
            para_match = arguments.get("paragraph_match", "")
            para_idx = arguments.get("paragraph_index")
            run_match = arguments.get("run_text_match", "")

            targets = []
            if para_idx is not None:
                targets = [doc.paragraphs[para_idx]]
            elif para_match:
                for p in doc.paragraphs:
                    if para_match.lower() in p.text.lower():
                        targets.append(p)

            for p in targets:
                for run in p.runs:
                    if not run_match or run_match.lower() in run.text.lower():
                        _set_run_formatting(
                            run,
                            bold=arguments.get("bold"),
                            italic=arguments.get("italic"),
                            underline=arguments.get("underline"),
                            strike=arguments.get("strike"),
                            font_size=arguments.get("font_size"),
                            font_name=arguments.get("font_name"),
                            color_hex=arguments.get("color_hex"),
                            all_caps=arguments.get("all_caps"),
                            small_caps=arguments.get("small_caps"),
                        )
            _save()
            return ok(f"Run formatting applied to {len(targets)} paragraph(s).")

        # ====================================================================
        # Headings
        # ====================================================================

        elif name == "add_heading":
            doc = _require_doc()
            text = arguments["text"]
            level = arguments.get("level", 1)
            after_heading = arguments.get("after_heading")

            if after_heading:
                ref = _find_heading_para(doc, after_heading)
                if ref:
                    style_name = f"Heading {level}"
                    style_id = _get_style_id(doc, style_name) or f"Heading{level}"
                    _insert_paragraph_after(ref, text, style_id)
                    _save()
                    return ok(f"Heading '{text}' (H{level}) inserted after '{after_heading}'.")

            p = doc.add_heading(text, level=level)
            _save()
            return ok(f"Heading '{text}' (H{level}) added.")

        # ====================================================================
        # Tables
        # ====================================================================

        elif name == "insert_table":
            doc = _require_doc()
            rows = arguments["rows"]
            cols = arguments["cols"]
            data = arguments.get("data", [])
            style = arguments.get("style", "Table Grid")
            header_row = arguments.get("header_row", False)
            col_widths = arguments.get("col_widths", [])

            tbl = doc.add_table(rows=rows, cols=cols)
            try:
                tbl.style = doc.styles[style]
            except Exception:
                pass

            for ri, row_data in enumerate(data):
                if ri >= rows:
                    break
                for ci, cell_val in enumerate(row_data):
                    if ci >= cols:
                        break
                    cell = tbl.cell(ri, ci)
                    if isinstance(cell_val, dict):
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(str(cell_val.get("text", "")))
                        _set_run_formatting(
                            run,
                            bold=cell_val.get("bold"),
                            italic=cell_val.get("italic"),
                            color_hex=cell_val.get("color_hex"),
                            font_size=cell_val.get("font_size"),
                        )
                        if cell_val.get("bg_color"):
                            _set_cell_background(cell, cell_val["bg_color"])
                    else:
                        cell.text = str(cell_val)

                    if header_row and ri == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

            if col_widths:
                for ci, width_cm in enumerate(col_widths):
                    if ci < cols:
                        for row in tbl.rows:
                            row.cells[ci].width = Cm(width_cm)

            _save()
            return ok(f"Table {rows}x{cols} inserted.")

        elif name == "edit_table_cell":
            doc = _require_doc()
            ti = arguments["table_index"]
            row = arguments["row"]
            col = arguments["col"]
            text = arguments["text"]
            cell = doc.tables[ti].cell(row, col)
            cell.text = text
            if arguments.get("bg_color"):
                _set_cell_background(cell, arguments["bg_color"])
            align_str = arguments.get("alignment")
            for para in cell.paragraphs:
                if align_str:
                    para.alignment = ALIGN_MAP.get(align_str.upper(), WD_ALIGN_PARAGRAPH.LEFT)
                for run in para.runs:
                    _set_run_formatting(
                        run,
                        bold=arguments.get("bold"),
                        italic=arguments.get("italic"),
                        font_size=arguments.get("font_size"),
                        color_hex=arguments.get("color_hex"),
                    )
            _save()
            return ok(f"Cell [{row},{col}] in table {ti} updated.")

        elif name == "add_table_row":
            doc = _require_doc()
            ti = arguments["table_index"]
            data = arguments.get("data", [])
            tbl = doc.tables[ti]
            row = tbl.add_row()
            for ci, cell_text in enumerate(data):
                if ci < len(row.cells):
                    row.cells[ci].text = str(cell_text)
            _save()
            return ok(f"Row added to table {ti}.")

        elif name == "merge_table_cells":
            doc = _require_doc()
            ti = arguments["table_index"]
            tbl = doc.tables[ti]
            direction = arguments["direction"]
            if direction == "horizontal":
                row = arguments["row"]
                start_col = arguments["start_col"]
                end_col = arguments["end_col"]
                _merge_cells_horizontal(tbl, row, start_col, end_col)
            else:
                col = arguments["col"]
                start_row = arguments["start_row"]
                end_row = arguments["end_row"]
                _merge_cells_vertical(tbl, col, start_row, end_row)
            _save()
            return ok(f"Cells merged ({direction}) in table {ti}.")

        elif name == "format_table_cell":
            doc = _require_doc()
            ti = arguments["table_index"]
            row = arguments["row"]
            col = arguments["col"]
            cell = doc.tables[ti].cell(row, col)
            if arguments.get("bg_color"):
                _set_cell_background(cell, arguments["bg_color"])
            border_args = {}
            for side in ["top", "bottom", "left", "right"]:
                val = arguments.get(f"border_{side}")
                if val:
                    border_args[side] = val
            if border_args:
                _set_cell_border(
                    cell,
                    color=arguments.get("border_color", "000000"),
                    size=arguments.get("border_size", 4),
                    **border_args
                )
            _save()
            return ok(f"Cell [{row},{col}] in table {ti} formatted.")

        elif name == "delete_table_row":
            doc = _require_doc()
            ti = arguments["table_index"]
            row_idx = arguments["row"]
            tbl = doc.tables[ti]
            row = tbl.rows[row_idx]
            row._tr.getparent().remove(row._tr)
            _save()
            return ok(f"Row {row_idx} deleted from table {ti}.")

        elif name == "delete_table":
            doc = _require_doc()
            ti = arguments["table_index"]
            tbl = doc.tables[ti]
            tbl._tbl.getparent().remove(tbl._tbl)
            _save()
            return ok(f"Table {ti} deleted.")

        # ====================================================================
        # Headers & Footers
        # ====================================================================

        elif name == "set_header":
            doc = _require_doc()
            text = arguments.get("text", "")
            si = arguments.get("section_index", 0)
            which = arguments.get("which", "header")
            alignment = arguments.get("alignment", "LEFT")
            section = doc.sections[si]
            hf = _get_or_create_hdrftr(section, which)
            # Clear existing content
            for p in hf.paragraphs:
                for run in p.runs:
                    run.text = ""
            para = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
            para.clear()
            para.alignment = ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run(text)
            _set_run_formatting(
                run,
                bold=arguments.get("bold"),
                italic=arguments.get("italic"),
                font_size=arguments.get("font_size"),
                font_name=arguments.get("font_name"),
            )
            _save()
            return ok(f"Header ({which}) set for section {si}.")

        elif name == "set_footer":
            doc = _require_doc()
            text = arguments.get("text", "")
            si = arguments.get("section_index", 0)
            which = arguments.get("which", "footer")
            alignment = arguments.get("alignment", "CENTER")
            add_page_number = arguments.get("add_page_number", False)
            page_number_format = arguments.get("page_number_format", "PAGE_NUMBER")
            section = doc.sections[si]
            hf = _get_or_create_hdrftr(section, which)
            para = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
            para.clear()
            para.alignment = ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)

            if text:
                run = para.add_run(text + ("  " if add_page_number else ""))
                _set_run_formatting(
                    run,
                    bold=arguments.get("bold"),
                    font_size=arguments.get("font_size"),
                    font_name=arguments.get("font_name"),
                )

            if add_page_number:
                _add_page_number_field(para, alignment=alignment, fmt=page_number_format)

            _save()
            return ok(f"Footer ({which}) set for section {si}.")

        elif name == "add_image_to_header":
            doc = _require_doc()
            image_path = arguments["image_path"]
            si = arguments.get("section_index", 0)
            width_cm = arguments.get("width_cm", 4.0)
            alignment = arguments.get("alignment", "LEFT")
            section = doc.sections[si]
            hf = _get_or_create_hdrftr(section, "header")
            para = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
            para.clear()
            para.alignment = ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run()
            run.add_picture(image_path, width=Cm(width_cm))
            _save()
            return ok(f"Image added to header of section {si}.")

        elif name == "clear_header":
            doc = _require_doc()
            si = arguments.get("section_index", 0)
            which = arguments.get("which", "header")
            section = doc.sections[si]
            hf = getattr(section, which)
            for p in hf.paragraphs:
                for run in p.runs:
                    run.text = ""
            _save()
            return ok(f"Header ({which}) cleared for section {si}.")

        elif name == "clear_footer":
            doc = _require_doc()
            si = arguments.get("section_index", 0)
            which = arguments.get("which", "footer")
            section = doc.sections[si]
            hf = getattr(section, which)
            for p in hf.paragraphs:
                for run in p.runs:
                    run.text = ""
            _save()
            return ok(f"Footer ({which}) cleared for section {si}.")

        # ====================================================================
        # Images
        # ====================================================================

        elif name == "insert_image":
            doc = _require_doc()
            image_path = arguments["image_path"]
            width_cm = arguments.get("width_cm")
            height_cm = arguments.get("height_cm")
            alignment = arguments.get("alignment", "LEFT")
            caption_text = arguments.get("caption")
            after_text = arguments.get("after_text")
            after_heading = arguments.get("after_heading")

            # Build the image paragraph
            img_para = doc.add_paragraph()
            img_para.alignment = ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
            run = img_para.add_run()

            kwargs = {}
            if width_cm:
                kwargs["width"] = Cm(width_cm)
            if height_cm:
                kwargs["height"] = Cm(height_cm)
            run.add_picture(image_path, **kwargs)

            # Reposition if needed (currently appended at end by python-docx API;
            # for precise positioning the XML element is moved)
            if after_text or after_heading:
                ref_para = None
                if after_heading:
                    ref_para = _find_heading_para(doc, after_heading)
                elif after_text:
                    for p in doc.paragraphs:
                        if after_text.lower() in p.text.lower():
                            ref_para = p
                            break
                if ref_para:
                    # Move the appended paragraph after ref_para
                    img_p_el = img_para._p
                    img_p_el.getparent().remove(img_p_el)
                    ref_para._p.addnext(img_p_el)

            if caption_text:
                cap_para = doc.add_paragraph(caption_text)
                try:
                    cap_para.style = doc.styles["Caption"]
                except Exception:
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Move caption after image
                cap_el = cap_para._p
                cap_el.getparent().remove(cap_el)
                img_para._p.addnext(cap_el)

            _save()
            return ok(f"Image '{image_path}' inserted.")

        # ====================================================================
        # Page layout
        # ====================================================================

        elif name == "set_page_margins":
            doc = _require_doc()
            si = arguments.get("section_index")
            sections = [doc.sections[si]] if si is not None else doc.sections
            for section in sections:
                if "top" in arguments:
                    section.top_margin = Cm(arguments["top"])
                if "bottom" in arguments:
                    section.bottom_margin = Cm(arguments["bottom"])
                if "left" in arguments:
                    section.left_margin = Cm(arguments["left"])
                if "right" in arguments:
                    section.right_margin = Cm(arguments["right"])
            _save()
            return ok("Page margins updated.")

        elif name == "set_page_orientation":
            doc = _require_doc()
            si = arguments.get("section_index", 0)
            orientation = arguments["orientation"].lower()
            section = doc.sections[si]
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width/height if not already landscape
                if section.page_width < section.page_height:
                    section.page_width, section.page_height = (
                        section.page_height, section.page_width
                    )
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                if section.page_width > section.page_height:
                    section.page_width, section.page_height = (
                        section.page_height, section.page_width
                    )
            _save()
            return ok(f"Section {si} orientation set to {orientation}.")

        elif name == "set_page_size":
            doc = _require_doc()
            si = arguments.get("section_index", 0)
            section = doc.sections[si]
            preset = arguments.get("preset")
            # Dimensions in EMU (1 cm = 360000 EMU)
            presets = {
                "A3": (Cm(29.7), Cm(42.0)),
                "A4": (Cm(21.0), Cm(29.7)),
                "A5": (Cm(14.8), Cm(21.0)),
                "Letter": (Cm(21.59), Cm(27.94)),
                "Legal": (Cm(21.59), Cm(35.56)),
            }
            if preset and preset in presets:
                w, h = presets[preset]
            elif "width_cm" in arguments and "height_cm" in arguments:
                w = Cm(arguments["width_cm"])
                h = Cm(arguments["height_cm"])
            else:
                return ok("Specify 'preset' or both 'width_cm' and 'height_cm'.")
            section.page_width = w
            section.page_height = h
            _save()
            size_label = preset or f"{arguments.get('width_cm')}x{arguments.get('height_cm')} cm"
            return ok(f"Page size set to {size_label}.")

        elif name == "add_section_break":
            doc = _require_doc()
            break_type = arguments.get("break_type", "new_page")
            break_map = {
                "new_page": WD_SECTION.NEW_PAGE,
                "continuous": WD_SECTION.CONTINUOUS,
                "even_page": WD_SECTION.EVEN_PAGE,
                "odd_page": WD_SECTION.ODD_PAGE,
            }
            after_text = arguments.get("after_text")
            ref_para = None
            if after_text:
                for p in doc.paragraphs:
                    if after_text.lower() in p.text.lower():
                        ref_para = p
                        break
            if ref_para:
                # Add section break via sectPr in the paragraph
                pPr = ref_para._p.get_or_add_pPr()
                sectPr = OxmlElement("w:sectPr")
                pgSz = OxmlElement("w:type")
                pgSz.set(qn("w:val"), list(break_map.keys())[
                    list(break_map.values()).index(break_map[break_type])
                ])
                sectPr.append(pgSz)
                pPr.append(sectPr)
            else:
                doc.add_section(break_map.get(break_type, WD_SECTION.NEW_PAGE))
            _save()
            return ok(f"Section break ({break_type}) inserted.")

        elif name == "add_page_break":
            doc = _require_doc()
            after_text = arguments.get("after_text")
            at_index = arguments.get("at_index")
            at_end = arguments.get("at_end", False)

            ref = None
            if after_text:
                for p in doc.paragraphs:
                    if after_text.lower() in p.text.lower():
                        ref = p
                        break
            elif at_index is not None:
                ref = doc.paragraphs[at_index]

            if ref:
                run = ref.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                doc.add_page_break()
            _save()
            return ok("Page break inserted.")

        elif name == "set_columns":
            doc = _require_doc()
            si = arguments.get("section_index", 0)
            num_cols = arguments["num_cols"]
            spacing_cm = arguments.get("spacing_cm", 1.25)
            _set_section_columns(doc.sections[si], num_cols, spacing_cm)
            _save()
            return ok(f"Section {si} set to {num_cols} column(s).")

        # ====================================================================
        # Styles
        # ====================================================================

        elif name == "create_style":
            doc = _require_doc()
            style = _create_paragraph_style(
                doc,
                style_name=arguments["style_name"],
                base_style=arguments.get("base_style", "Normal"),
                font_name=arguments.get("font_name"),
                font_size=arguments.get("font_size"),
                bold=arguments.get("bold"),
                italic=arguments.get("italic"),
                color_hex=arguments.get("color_hex"),
                alignment=arguments.get("alignment"),
                space_before=arguments.get("space_before"),
                space_after=arguments.get("space_after"),
            )
            _save()
            return ok(f"Style '{style.name}' created/updated.")

        # ====================================================================
        # TOC
        # ====================================================================

        elif name == "insert_toc":
            doc = _require_doc()
            title = arguments.get("title", "Table of Contents")
            max_level = arguments.get("max_level", 3)
            after_heading = arguments.get("after_heading")
            at_end = arguments.get("at_end", False)
            _insert_toc(doc, title=title, max_level=max_level)
            _save()
            return ok(
                f"TOC inserted (levels 1-{max_level}). "
                "Open in Word and press Ctrl+A then F9 to render the TOC."
            )

        # ====================================================================
        # Hyperlinks
        # ====================================================================

        elif name == "add_hyperlink":
            doc = _require_doc()
            para_match = arguments.get("paragraph_match", "")
            para_idx = arguments.get("paragraph_index")
            link_text = arguments["text"]
            url = arguments["url"]
            color_hex = arguments.get("color_hex", "0563C1")
            underline = arguments.get("underline", True)

            target = None
            if para_idx is not None:
                target = doc.paragraphs[para_idx]
            elif para_match:
                for p in doc.paragraphs:
                    if para_match.lower() in p.text.lower():
                        target = p
                        break
            if target is None:
                return ok("Paragraph not found.")
            _add_hyperlink(target, link_text, url, color_hex=color_hex, underline=underline)
            _save()
            return ok(f"Hyperlink '{link_text}' -> '{url}' added.")

        # ====================================================================
        # Comments
        # ====================================================================

        elif name == "add_comment":
            doc = _require_doc()
            match_text = arguments["match_text"]
            comment_text = arguments["comment"]
            author = arguments.get("author", "Claude")
            initials = arguments.get("initials", "AI")
            for p in doc.paragraphs:
                if match_text.lower() in p.text.lower():
                    cid = _add_native_comment(doc, p, comment_text,
                                              author=author, initials=initials)
                    _save()
                    return ok(f"Comment (id={cid}) added to paragraph containing '{match_text}'.")
            return ok(f"Text '{match_text}' not found.")

        # ====================================================================
        # Bookmarks
        # ====================================================================

        elif name == "add_bookmark":
            doc = _require_doc()
            para_idx = arguments.get("paragraph_index")
            match_text = arguments.get("match_text", "")
            bookmark_name = arguments["bookmark_name"]
            bookmark_id = arguments.get("bookmark_id", 1)

            target = None
            if para_idx is not None:
                target = doc.paragraphs[para_idx]
            elif match_text:
                for p in doc.paragraphs:
                    if match_text.lower() in p.text.lower():
                        target = p
                        break
            if target is None:
                return ok("Paragraph not found.")
            _add_bookmark(target, bookmark_name, bookmark_id)
            _save()
            return ok(f"Bookmark '{bookmark_name}' added.")

        # ====================================================================
        # Document properties
        # ====================================================================

        elif name == "set_document_properties":
            doc = _require_doc()
            cp = doc.core_properties
            if "title" in arguments:
                cp.title = arguments["title"]
            if "author" in arguments:
                cp.author = arguments["author"]
            if "subject" in arguments:
                cp.subject = arguments["subject"]
            if "keywords" in arguments:
                cp.keywords = arguments["keywords"]
            if "description" in arguments:
                cp.description = arguments["description"]
            if "category" in arguments:
                cp.category = arguments["category"]
            _save()
            return ok("Document properties updated.")

        # ====================================================================
        # Advanced XML
        # ====================================================================

        elif name == "apply_xml_patch":
            doc = _require_doc()
            idx = arguments["paragraph_index"]
            xml_content = arguments["xml_content"]
            old_p = doc.paragraphs[idx]._p
            new_p = etree.fromstring(xml_content.encode())
            old_p.getparent().replace(old_p, new_p)
            _save()
            return ok(f"XML patch applied to paragraph {idx}.")

        # ====================================================================
        # Batch document builder
        # ====================================================================

        elif name == "build_document":
            path = arguments["path"]
            template = arguments.get("template")
            doc = Document(template) if template else Document()

            # Default font/size
            if arguments.get("default_font") or arguments.get("default_font_size"):
                normal = doc.styles["Normal"]
                if arguments.get("default_font"):
                    normal.font.name = arguments["default_font"]
                if arguments.get("default_font_size"):
                    normal.font.size = Pt(arguments["default_font_size"])

            # Document properties
            cp = doc.core_properties
            for prop in ["title", "author", "subject"]:
                if arguments.get(prop):
                    setattr(cp, prop, arguments[prop])

            # Margins
            if arguments.get("margins"):
                m = arguments["margins"]
                for section in doc.sections:
                    if "top" in m:
                        section.top_margin = Cm(m["top"])
                    if "bottom" in m:
                        section.bottom_margin = Cm(m["bottom"])
                    if "left" in m:
                        section.left_margin = Cm(m["left"])
                    if "right" in m:
                        section.right_margin = Cm(m["right"])

            def apply_run_fmt(run, data: dict):
                _set_run_formatting(
                    run,
                    bold=data.get("bold"),
                    italic=data.get("italic"),
                    underline=data.get("underline"),
                    strike=data.get("strike"),
                    font_size=data.get("font_size"),
                    font_name=data.get("font_name"),
                    color_hex=data.get("color_hex"),
                )

            def apply_para_format(para, el: dict):
                if el.get("alignment"):
                    para.alignment = ALIGN_MAP.get(el["alignment"].upper(), WD_ALIGN_PARAGRAPH.LEFT)
                pf = para.paragraph_format
                if el.get("space_before") is not None:
                    pf.space_before = Pt(el["space_before"])
                if el.get("space_after") is not None:
                    pf.space_after = Pt(el["space_after"])

            def fill_paragraph(para, el: dict):
                """Fill paragraph with text or runs (with optional hyperlinks)."""
                if el.get("runs"):
                    for rd in el["runs"]:
                        if rd.get("url"):
                            _add_hyperlink(para, rd.get("text", ""), rd["url"],
                                           color_hex=rd.get("color_hex", "0563C1"))
                        else:
                            run = para.add_run(rd.get("text", ""))
                            apply_run_fmt(run, rd)
                elif el.get("text"):
                    para.add_run(el["text"])

            added = 0
            for el in arguments.get("elements", []):
                el_type = el.get("type", "paragraph")

                if el_type == "page_break":
                    doc.add_page_break()

                elif el_type == "section_break":
                    break_map = {
                        "new_page": WD_SECTION.NEW_PAGE,
                        "continuous": WD_SECTION.CONTINUOUS,
                        "even_page": WD_SECTION.EVEN_PAGE,
                        "odd_page": WD_SECTION.ODD_PAGE,
                    }
                    bt = el.get("break_type", "new_page")
                    doc.add_section(break_map.get(bt, WD_SECTION.NEW_PAGE))

                elif el_type == "toc":
                    _insert_toc(doc,
                                title=el.get("toc_title", "Table of Contents"),
                                max_level=el.get("toc_max_level", 3))

                elif el_type == "heading":
                    level = el.get("level", 1)
                    para = doc.add_heading("", level=level)
                    fill_paragraph(para, el)
                    apply_para_format(para, el)

                elif el_type in ("paragraph", "list"):
                    list_type = el.get("list_type")
                    list_level = el.get("list_level", 0)
                    style_name = el.get("style")
                    if list_type == "bullet":
                        style_name = style_name or (
                            "List Bullet" if list_level == 0
                            else f"List Bullet {list_level + 1}"
                        )
                    elif list_type == "numbered":
                        style_name = style_name or (
                            "List Number" if list_level == 0
                            else f"List Number {list_level + 1}"
                        )
                    else:
                        style_name = style_name or "Normal"
                    try:
                        para = doc.add_paragraph(style=style_name)
                    except Exception:
                        para = doc.add_paragraph()
                    fill_paragraph(para, el)
                    apply_para_format(para, el)

                elif el_type == "image":
                    image_path = el.get("image_path", "")
                    if image_path and os.path.exists(image_path):
                        img_para = doc.add_paragraph()
                        alignment = el.get("alignment", "LEFT")
                        img_para.alignment = ALIGN_MAP.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
                        run = img_para.add_run()
                        kwargs = {}
                        if el.get("width_cm"):
                            kwargs["width"] = Cm(el["width_cm"])
                        run.add_picture(image_path, **kwargs)
                        if el.get("caption"):
                            cap = doc.add_paragraph(el["caption"])
                            try:
                                cap.style = doc.styles["Caption"]
                            except Exception:
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif el_type == "table":
                    data = el.get("data", [])
                    if not data:
                        added += 1
                        continue
                    rows = len(data)
                    cols = max(len(r) for r in data) if data else 1
                    table_style = el.get("table_style", "Table Grid")
                    tbl = doc.add_table(rows=rows, cols=cols)
                    try:
                        tbl.style = doc.styles[table_style]
                    except Exception:
                        pass
                    col_widths = el.get("col_widths", [])
                    header_row = el.get("header_row", False)
                    for ri, row_data in enumerate(data):
                        for ci, cell_val in enumerate(row_data):
                            if ci >= cols:
                                break
                            cell = tbl.cell(ri, ci)
                            if isinstance(cell_val, dict):
                                cell.text = ""
                                run = cell.paragraphs[0].add_run(
                                    str(cell_val.get("text", ""))
                                )
                                apply_run_fmt(run, cell_val)
                                if cell_val.get("bg_color"):
                                    _set_cell_background(cell, cell_val["bg_color"])
                            else:
                                cell.text = str(cell_val)
                            if header_row and ri == 0:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.bold = True
                    if col_widths:
                        for ci, width_cm in enumerate(col_widths):
                            if ci < cols:
                                for row in tbl.rows:
                                    row.cells[ci].width = Cm(width_cm)

                elif el_type == "hyperlink":
                    para = doc.add_paragraph()
                    style_name = el.get("style", "Normal")
                    try:
                        para.style = doc.styles[style_name]
                    except Exception:
                        pass
                    apply_para_format(para, el)
                    _add_hyperlink(
                        para,
                        text=el.get("text", el.get("url", "")),
                        url=el.get("url", ""),
                        color_hex=el.get("color_hex", "0563C1"),
                    )

                added += 1

            # Global header
            if arguments.get("header_text"):
                hdr = doc.sections[0].header
                p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
                p.clear()
                p.add_run(arguments["header_text"])

            # Global footer
            if arguments.get("footer_text") or arguments.get("footer_page_numbers"):
                ftr = doc.sections[0].footer
                p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if arguments.get("footer_text"):
                    p.add_run(arguments["footer_text"])
                if arguments.get("footer_page_numbers"):
                    _add_page_number_field(p, alignment="CENTER", fmt="PAGE_OF_PAGES")

            doc.save(path)
            _current_doc_path = path
            _doc = Document(path)
            return ok(
                f"Document built: {path}\n"
                f"Elements: {added} | "
                f"Paragraphs: {len(_doc.paragraphs)} | "
                f"Tables: {len(_doc.tables)}"
            )

        else:
            return ok(f"Unknown tool: {name}")

    except Exception as exc:
        import traceback
        return ok(f"ERROR: {exc}\n{traceback.format_exc()}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

async def main():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream, write_stream,
            server.create_initialization_options()
        )


if __name__ == "__main__":
    asyncio.run(main())
